"""
Appends the Fairness Index section (Section 9) to the existing report .docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Helpers (same as report generator) ──────────────────────────────────────
TEAL  = RGBColor(0x00, 0x96, 0x88)
DARK  = RGBColor(0x1A, 0x27, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED   = RGBColor(0xEF, 0x44, 0x44)
GREEN = RGBColor(0x10, 0xB9, 0x81)
GRAY  = RGBColor(0x6B, 0x72, 0x80)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if color:
            run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg="006064"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, "F0F4F5")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ─── Scores ───────────────────────────────────────────────────────────────────
DIMENSIONS = [
    {
        "name": "Participation Equity",
        "score": 97.7,
        "weight": "20%",
        "rating": "Strong",
        "rating_color": "10B981",
        "key_metric": "96.9% seniority-verified; 97.5% active; 100% submission acceptance",
        "interpretation": (
            "Every employee who submitted a request was accepted into the system without "
            "gatekeeping. The 96.9% verification rate means that nearly all seniority tie-breakers "
            "are based on confirmed, audited dates rather than self-reported estimates. The 5 "
            "unverified accounts represent the only structural gap in this dimension."
        ),
        "improvement": "Resolve the 5 pending verifications before deliberations open to achieve a perfect score.",
    },
    {
        "name": "Priority Access Equity",
        "score": 98.7,
        "weight": "25%",
        "rating": "Strong",
        "rating_color": "10B981",
        "key_metric": "156/158 P1 requests (98.7%) fall within the 8-cap on ≥1 date; all seniority bands = 100% access",
        "interpretation": (
            "This is the most heavily weighted dimension because it directly measures whether the "
            "system's core promise — that your top-priority request has a realistic path to approval — "
            "holds across all seniority levels. Junior nurses (0–4 years) achieve the same 100% cap "
            "access rate as 20+ year veterans on their P1 requests. This is a structural achievement: "
            "the 8-person cap is large enough relative to shift size that even the most junior "
            "employees can access it on their chosen dates. Only 2 P1 requests are entirely outside "
            "the cap on every date they requested — both on the AM shift during the peak "
            "July–October window."
        ),
        "improvement": "Address the 2 fully-displaced P1 requests through direct admin conversation about alternative dates.",
    },
    {
        "name": "Shift Parity",
        "score": 52.0,
        "weight": "15%",
        "rating": "Needs Attention",
        "rating_color": "EF4444",
        "key_metric": "AM cap utilization 72% vs NOC 53.1% vs PM 17.1%; AM oversubscribed 33.8% of days",
        "interpretation": (
            "This is the lowest-scoring dimension and the most structurally significant finding in "
            "the Fairness Index. The AM shift bears a disproportionate deliberation burden: 73 of its "
            "216 active days (33.8%) are oversubscribed, compared to 21 of 243 NOC days (8.6%) and "
            "zero PM days. The root cause is not unfairness in the rules — the same 8-person cap "
            "applies to all shifts — but a fundamental imbalance in shift size. With 80 AM employees "
            "and only 9 PM employees, the cap represents 10% of AM capacity and 89% of PM capacity. "
            "An AM nurse faces a structurally harder approval environment than a PM nurse, purely "
            "because of which shift they work. This is not a portal design failure; it is a workforce "
            "composition reality that the portal has made visible for the first time."
        ),
        "improvement": (
            "Consider a shift-proportional cap formula (e.g., 10% of shift headcount, rounded up) "
            "rather than a flat 8. For AM: 8 (unchanged). For NOC: 7. For PM: 1–2. "
            "This would equalize the structural burden across shifts and is a policy decision for "
            "unit leadership, not a technical change."
        ),
    },
    {
        "name": "Process Transparency",
        "score": 79.7,
        "weight": "15%",
        "rating": "Moderate",
        "rating_color": "F59E0B",
        "key_metric": "96.9% verified; 33.3% added comments; 13.7% used priority re-ranking",
        "interpretation": (
            "Process transparency measures whether employees can meaningfully signal their preferences "
            "and whether those signals are captured with integrity. The verification rate is strong. "
            "Comment usage at 33.3% is healthy for a first-cycle deployment — one in three employees "
            "chose to add context to their requests, which is a sign of trust in the system. Priority "
            "re-ranking (13.7% of requests) indicates that employees are actively managing their "
            "submissions rather than submitting and forgetting. The score is held at Moderate rather "
            "than Strong because 3.1% of accounts remain unverified, and because comment and "
            "re-ranking adoption has room to grow in future cycles."
        ),
        "improvement": (
            "In Cycle 2, add an in-portal prompt reminding employees to review and confirm their "
            "priority rankings 7 days before the submission deadline. This would likely increase "
            "re-ranking engagement and push this score into the Strong range."
        ),
    },
    {
        "name": "Demand Concentration",
        "score": 69.4,
        "weight": "15%",
        "rating": "Moderate",
        "rating_color": "F59E0B",
        "key_metric": "Gini coefficient 0.306; 82 employees (50.6%) submitted 1–2 requests; 7 employees submitted 5+",
        "interpretation": (
            "Demand concentration measures how evenly distributed requests are across the employee "
            "population. A Gini coefficient of 0.306 indicates moderate inequality — the distribution "
            "is not perfectly equal (where every employee submits the same number of requests), but "
            "it is far from highly concentrated. The majority of employees (82 of 157, or 52.2%) "
            "submitted 1–2 active vacation requests, while 7 employees submitted 5 or more. The "
            "portal's policy allows up to 10 requests, and the system does not penalize high-volume "
            "submitters — their additional requests simply receive lower working priority. This is "
            "the intended design: more requests mean more options for the employee, but the "
            "first-priority request is what matters most."
        ),
        "improvement": (
            "Monitor whether high-volume submitters (5+ requests) are consuming a disproportionate "
            "share of approved slots in Cycle 2. If so, consider a soft limit of 5 active requests "
            "per employee to reduce administrative review volume."
        ),
    },
    {
        "name": "Ceiling Equity",
        "score": 80.3,
        "weight": "10%",
        "rating": "Strong",
        "rating_color": "10B981",
        "key_metric": "AM cap = 10% of shift; NOC = 11.3%; PM = 88.9% (structurally unconstrained)",
        "interpretation": (
            "Ceiling equity measures whether the 8-person cap is appropriately calibrated for each "
            "shift's headcount. For AM and NOC, the cap represents 10–11% of shift size, which falls "
            "within the 8–15% ideal range for a critical care unit. For PM, the cap is effectively "
            "unconstrained: with only 9 PM employees, an 8-person cap means 88.9% of the shift could "
            "theoretically be off on the same day. In practice, PM has zero oversubscribed days, so "
            "this does not create a real operational problem in Cycle 1. However, it does mean that "
            "PM employees face a structurally easier approval environment, which is the same "
            "structural imbalance identified in Shift Parity."
        ),
        "improvement": (
            "The PM cap is a non-issue operationally in Cycle 1, but should be revisited if PM "
            "headcount grows. A proportional cap would bring this dimension to a perfect score."
        ),
    },
]

COMPOSITE = 82.4

# ─── Load and append to existing report ───────────────────────────────────────
doc_path = "/home/ubuntu/vnc-icu-portal/exports/VNC_ICU_Portal_Report_May2026.docx"
doc = Document(doc_path)

# Page break before new section
doc.add_page_break()

# ── SECTION 9 HEADER ──────────────────────────────────────────────────────────
add_heading(doc, "9. Fairness Index — Cycle 1 Assessment", 1, TEAL)
add_para(doc,
    "The VNC ICU Portal Fairness Index is a purpose-built, multi-dimensional scoring framework "
    "designed to evaluate the equity of the vacation request process from six distinct angles. "
    "Each dimension is scored on a 0–100 scale using live data from the portal database and the "
    "working priority CSV. The composite score is a weighted average reflecting the relative "
    "importance of each dimension to the unit's fairness mission. This index is intended to be "
    "recomputed at the close of each deliberation cycle, creating a year-over-year benchmark "
    "for process improvement."
)
doc.add_paragraph()

# ── COMPOSITE SCORE BOX ───────────────────────────────────────────────────────
add_heading(doc, "Composite Fairness Score: 82.4 / 100 — Strong", 2, GREEN)
add_para(doc,
    "The portal earns a composite score of 82.4 out of 100 in its inaugural cycle — a Strong "
    "rating. This reflects a system that is fundamentally equitable in its access rules and "
    "participation design, with one significant structural gap (Shift Parity) that is rooted in "
    "workforce composition rather than system design, and two dimensions (Process Transparency "
    "and Demand Concentration) that are Moderate and expected to improve as staff become more "
    "familiar with the portal in future cycles."
)
doc.add_paragraph()

# ── RADAR CHART ───────────────────────────────────────────────────────────────
add_para(doc, "Figure 1: Fairness Index Radar Chart — Cycle 1", bold=True, size=10, color=GRAY)
doc.add_picture(
    "/home/ubuntu/vnc-icu-portal/exports/fairness_radar.png",
    width=Inches(4.5)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── SCORE SUMMARY TABLE ───────────────────────────────────────────────────────
add_heading(doc, "Score Summary", 2, DARK)
make_table(doc,
    ["Dimension", "Score", "Weight", "Rating", "Key Metric"],
    [
        [
            d["name"],
            f"{d['score']}/100",
            d["weight"],
            d["rating"],
            d["key_metric"],
        ]
        for d in DIMENSIONS
    ] + [
        ["COMPOSITE (weighted)", f"{COMPOSITE}/100", "100%", "Strong", "Weighted average of all 6 dimensions"],
    ],
    col_widths=[1.8, 0.7, 0.6, 0.9, 2.8]
)
doc.add_paragraph()

# ── RATING SCALE ─────────────────────────────────────────────────────────────
add_heading(doc, "Rating Scale", 2, DARK)
make_table(doc,
    ["Score Range", "Rating", "Interpretation"],
    [
        ["80–100", "Strong",           "The process is performing well on this dimension. Maintain and monitor."],
        ["60–79",  "Moderate",         "Acceptable performance with clear improvement opportunities. Target in next cycle."],
        ["40–59",  "Needs Attention",  "A structural gap exists. Requires policy or design intervention before Cycle 2."],
        ["0–39",   "Critical",         "Systemic inequity. Immediate corrective action required."],
    ],
    col_widths=[1.2, 1.2, 4.4]
)
doc.add_paragraph()

# ── DIMENSION NARRATIVES ──────────────────────────────────────────────────────
add_heading(doc, "Dimension-by-Dimension Analysis", 2, DARK)

for i, d in enumerate(DIMENSIONS, 1):
    # Dimension heading with score
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"D{i}. {d['name']}  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = TEAL
    r2 = p.add_run(f"{d['score']}/100 — {d['rating']}")
    r2.bold = True
    r2.font.size = Pt(11)
    # Color by rating
    if d["rating"] == "Strong":
        r2.font.color.rgb = GREEN
    elif d["rating"] == "Moderate":
        r2.font.color.rgb = AMBER
    else:
        r2.font.color.rgb = RED

    # Key metric
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run("Key metric: ")
    r3.bold = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = GRAY
    r4 = p2.add_run(d["key_metric"])
    r4.font.size = Pt(10)
    r4.font.color.rgb = GRAY
    r4.italic = True

    # Interpretation
    add_para(doc, d["interpretation"], size=10.5, space_before=2, space_after=2)

    # Improvement
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(6)
    r5 = p3.add_run("Recommended action: ")
    r5.bold = True
    r5.font.size = Pt(10)
    r5.font.color.rgb = TEAL
    r6 = p3.add_run(d["improvement"])
    r6.font.size = Pt(10)
    r6.font.color.rgb = DARK

doc.add_paragraph()

# ── METHODOLOGY NOTE ──────────────────────────────────────────────────────────
add_heading(doc, "Methodology Note", 2, DARK)
add_para(doc,
    "All scores are computed from live database queries and the working priority CSV generated "
    "on May 9, 2026. The framework draws on established fairness measurement concepts including "
    "the Gini coefficient (for demand concentration), proportional cap analysis (for ceiling "
    "equity), and participation completeness metrics (for process transparency). Dimension weights "
    "reflect the deliberation committee's stated priorities: Priority Access Equity (25%) is "
    "weighted highest because it directly measures whether the system's core promise to employees "
    "is kept. Participation Equity (20%) is second because a fair process must be accessible to "
    "all. The remaining four dimensions share equal weight at 10–15% each. Weights and scoring "
    "formulas are documented in the portal's scripts directory (scripts/fairness-index.py) and "
    "can be adjusted by the admin team for future cycles.",
    size=10, color=GRAY
)
add_para(doc,
    "This index is not a final verdict on the fairness of individual decisions — those remain "
    "the responsibility of the management team. It is a structural diagnostic: a way of seeing "
    "where the system is working well, where it is working adequately, and where it needs "
    "deliberate improvement before the next cycle begins.",
    size=10, color=GRAY, italic=True
)

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(doc_path)
print(f"Fairness Index appended and saved: {doc_path}")
